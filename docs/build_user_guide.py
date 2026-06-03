"""
Generate the SENDA User Guide .docx.

- Cover, Acknowledgement, About Us, Table of Contents
- One section per page (alternating image / text columns)
- "Becoming a Partner" walkthrough
- Team page at the end

Pages with no screenshot file are silently excluded ("coming soon").
Run:    python docs/build_user_guide.py
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SCREENSHOT_DIR = ROOT / "screenshots"
OUTPUT_PATH = ROOT / "Senda-User-Guide.docx"

# ---- Brand palette ------------------------------------------------------

BRAND_PRIMARY = RGBColor(0x16, 0x4E, 0x9E)
BRAND_ACCENT = RGBColor(0x1F, 0x95, 0x4E)
TEXT_DARK = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_MUTED = RGBColor(0x55, 0x55, 0x55)
TEXT_FAINT = RGBColor(0x7B, 0x83, 0x8E)
TIP_FILL = "F4F8FD"
SECTION_RULE = "DDE2EA"

# ---- Page data ---------------------------------------------------------

@dataclass
class PageDoc:
    slug: str
    title: str
    route: str
    purpose: str
    sections: List[str]                 # bullet list of what's on the page
    how_to: List[str] = field(default_factory=list)   # numbered steps
    benefits: List[str] = field(default_factory=list) # why this matters
    tip: str = ""                       # short italic tip line

PAGES: List[PageDoc] = [
    PageDoc(
        slug="landing",
        title="Landing Page",
        route="/",
        purpose="The public face of SENDA — the page anyone visiting sms.mifumolabs.com sees first. It explains what SENDA does, shows pricing, and invites visitors to sign up.",
        sections=[
            "Hero section with an animated phone mock-up showing real business SMS conversations.",
            "Features grid covering SMS Campaigns, WhatsApp Business, Contact Management, Analytics, Templates and Multi-language support.",
            "Pricing tiers (Lite, Standard, Pro) and a live SMS cost calculator so visitors can estimate spend.",
            "Trust indicators — businesses served, total messages delivered, delivery rate.",
            "Tutorial video modal that walks new users through the platform.",
        ],
        how_to=[
            "Click Get Started or Sign Up to open the registration page.",
            "Click Login if you already have an account.",
            "Click How to Use to watch the tutorial video.",
            "Use the language toggle in the top bar to switch between English and Kiswahili.",
        ],
        benefits=[
            "Visitors instantly understand what SENDA offers without needing to log in.",
            "Pricing is transparent — costs are visible before signing up.",
        ],
    ),
    PageDoc(
        slug="login",
        title="Login",
        route="/login",
        purpose="Where returning users sign in with their email and password to reach their dashboard.",
        sections=[
            "Email and Password fields with show/hide visibility toggle.",
            "Remember me checkbox for trusted devices.",
            "Forgot Password link for account recovery.",
            "Sign-up link for first-time visitors.",
            "Side panel (on desktop) showing key benefits of SENDA.",
        ],
        how_to=[
            "Type the email address you used when signing up.",
            "Enter your password.",
            "Tick Remember me if you are on your personal device.",
            "Click Sign In.",
            "If sign-in fails, click Forgot Password to reset.",
        ],
        tip="Never share your password. SENDA staff will never ask you for it.",
    ),
    PageDoc(
        slug="signup",
        title="Signup",
        route="/signup",
        purpose="Where new users create an account, verify their phone, and (optionally) submit basic business details.",
        sections=[
            "Personal details — First Name, Last Name, Email, Phone, Country.",
            "Business details — Company name.",
            "Password and Confirm Password with strength indicators.",
            "Terms of Service and Privacy Policy acceptance checkbox.",
            "Phone verification screen — 6-digit code entry with Resend option.",
        ],
        how_to=[
            "Fill in your personal and business details.",
            "Create a strong password (minimum 8 characters, mix letters and numbers).",
            "Read and accept the Terms of Service and Privacy Policy.",
            "Click Sign Up — you will receive a 6-digit verification code by SMS.",
            "Enter the code on the verification screen and click Verify.",
            "If the SMS does not arrive, click Resend or Contact Support.",
        ],
        benefits=[
            "Signup is free and takes less than two minutes.",
            "Your account is verified by phone number so messaging is secure from day one.",
        ],
    ),
    PageDoc(
        slug="forgot-password",
        title="Forgot Password",
        route="/forgot-password",
        purpose="A three-step recovery flow that uses SMS verification to safely reset a forgotten password.",
        sections=[
            "Step 1 — enter the phone number tied to your account.",
            "Step 2 — receive a 6-digit reset code by SMS and enter it.",
            "Step 3 — set a new password with show/hide toggles.",
        ],
        how_to=[
            "Enter the phone number registered with your account.",
            "Click Send Reset Code.",
            "Open the SMS, copy the 6-digit code, and paste it on the next screen.",
            "Set a new password and confirm it.",
            "Click Reset Password and sign in with the new credentials.",
        ],
        tip="If the code does not arrive within 60 seconds, click Resend. Make sure the phone has SMS reception.",
    ),
    PageDoc(
        slug="reset-password",
        title="Reset Password",
        route="/reset-password",
        purpose="The page that opens when a user clicks the secure reset link sent to their email or phone. It lets them choose a new password.",
        sections=[
            "Card with a key icon for visual confirmation.",
            "New Password and Confirm Password fields, each with a show/hide toggle.",
            "Password requirements helper (minimum 8 characters).",
            "Success state with a green checkmark once the password is changed.",
        ],
        how_to=[
            "Click the reset link in the SMS or email — this opens the page automatically.",
            "Type a new password that you have not used before.",
            "Re-type the same password in Confirm.",
            "Click Reset Password.",
            "When you see the success message, click Sign In to Account.",
        ],
    ),
    PageDoc(
        slug="sms-activation",
        title="SMS / Account Activation",
        route="/smsactivation",
        purpose="Verifies the phone number tied to a newly registered account so the platform knows the user owns it.",
        sections=[
            "Large 6-digit code input field.",
            "Resend SMS button with a countdown timer.",
            "Switch between SMS and Email verification methods.",
            "Success card with confirmation and a Go to Dashboard button.",
        ],
        how_to=[
            "Wait for the SMS code to arrive on the phone you registered with.",
            "Type the 6-digit code into the input field.",
            "Click Activate Account.",
            "Once activated, click Go to Dashboard to start using SENDA.",
        ],
        tip="If you registered with the wrong number, sign out and sign up again with the correct one — phone numbers cannot be changed during activation.",
    ),
    PageDoc(
        slug="dashboard",
        title="Home Dashboard",
        route="/dashboard",
        purpose="The home page after login. A single hub showing the most important numbers, quick actions, and recent activity so users always know where they stand.",
        sections=[
            "Welcome banner with the user's first name and top two highlight metrics.",
            "Four metric cards — Total Messages sent, Active Contacts, Current Credits, and Approved Sender IDs.",
            "Quick Actions card — one-tap shortcuts to Send Message, Create Campaign and Manage Contacts.",
            "Performance Overview chart showing message volume over time.",
            "Activity Feed and Recent Campaigns side by side.",
            "Sender IDs status table showing which sender names are approved, pending or awaiting payment.",
            "Getting Started wizard for brand-new users with guided setup steps.",
        ],
        how_to=[
            "Read your name and the highlight metrics at the top to confirm you're in the right account.",
            "Tap a metric card to drill into the detail page (e.g. tap Total Messages to open Send SMS history).",
            "Use Quick Actions to start a new task without hunting through the menu.",
            "Check Sender IDs near the bottom — you cannot send branded SMS until at least one sender ID is approved.",
        ],
        benefits=[
            "Everything important is one tap away — no need to dig through menus.",
            "New users see a Getting Started wizard so they aren't lost on first login.",
        ],
    ),
    PageDoc(
        slug="notifications",
        title="Notifications",
        route="/notifications",
        purpose="The central activity feed — every important event (message sent, campaign finished, contact imported, template approved) lands here.",
        sections=[
            "Search bar at the top.",
            "Filter chips — All, Messages, Campaigns, Contacts.",
            "Notifications grouped by date (Today, Yesterday, Earlier).",
            "Each row shows an icon, title, short description, time, and a dismiss button.",
            "Tapping a notification opens a detail sheet with full context.",
        ],
        how_to=[
            "Open Notifications from the main menu or the bell icon in the top bar.",
            "Use the search bar to look up a specific event by name.",
            "Tap a filter chip to narrow the list to one type of activity.",
            "Tap any notification to see the full detail.",
            "Dismiss notifications you no longer need by tapping the X.",
        ],
        tip="Notifications older than 30 days are automatically archived to keep the list clean.",
    ),
    PageDoc(
        slug="contacts",
        title="Contacts",
        route="/contacts",
        purpose="The customer database — every phone number you send messages to lives here. Contacts can be imported, organised with tags, and used for bulk sends.",
        sections=[
            "Header buttons — Import, Refresh, Export.",
            "Search bar and tag filter for narrowing the list.",
            "Contact table with Name, Phone, Email, Tags and Status badges.",
            "Bulk action bar that appears when you tick multiple rows (Delete, Add Tags, Send SMS/WhatsApp, Edit).",
            "Create Contact dialog with fields for Name, Phone, Email, Company, Department and Tags.",
        ],
        how_to=[
            "Add one contact: click Create Contact, fill in at least Name and Phone, click Save.",
            "Import many at once: click Import → choose 'From Phone' (on mobile) or 'CSV File' and follow the prompts.",
            "Organise with tags: open a contact, type a tag name (e.g. 'VIP', 'Nairobi', 'Customer'), press Enter.",
            "Bulk send: tick contacts using the checkboxes, then click 'Send SMS' or 'Send WhatsApp' from the action bar.",
            "Export: click Export to download all contacts as a CSV for backup or external use.",
        ],
        benefits=[
            "Contacts are the foundation of every campaign — no contacts, no bulk messaging.",
            "Tags let you build re-usable audiences (e.g. send a Friday promo only to contacts tagged 'VIP').",
            "Imported once, contacts are available across SMS, WhatsApp and AI Copilots.",
        ],
        tip="Always store phone numbers in international format (e.g. +2557…) — local formats may fail in some networks.",
    ),
    PageDoc(
        slug="send-hub",
        title="Send Hub",
        route="/send",
        purpose="A simple channel chooser that asks 'How do you want to send?' so first-time users do not have to guess where to start.",
        sections=[
            "Two large channel cards side by side — Send an SMS and Send on WhatsApp.",
            "Each card has a short description of when to use it.",
            "Below the cards, a secondary Create Campaign call-to-action for bulk sends.",
        ],
        how_to=[
            "Tap Continue with SMS to compose a one-off SMS or a bulk SMS blast.",
            "Tap Send on WhatsApp to send a WhatsApp template or media message.",
            "Tap Create Campaign for scheduled, large-audience sends with tracking.",
        ],
        tip="Use SMS when you need to reach every phone (works without data). Use WhatsApp when you want rich media, links and two-way chat.",
    ),
    PageDoc(
        slug="send-sms",
        title="Send SMS",
        route="/sms/send",
        purpose="The main SMS composition screen — write a message, pick recipients, and either send right away or schedule for later. It is the most-used page on the platform.",
        sections=[
            "Sender name dropdown — only approved sender IDs are shown.",
            "Three sending modes — Quick (one number), Bulk (paste many numbers), Segments (saved audience).",
            "Message composer with live character counter showing SMS segments and total cost.",
            "Preview that mimics how the SMS will look on a recipient's phone.",
            "Recipient input — single number, bulk paste, tag group, or saved segment.",
            "Schedule toggle — Send Now or pick a future date and time.",
            "Cost summary and final Send button with a progress indicator while sending.",
        ],
        how_to=[
            "Pick the sender name your recipients should see (e.g. 'YOURBRAND').",
            "Choose the mode — Quick, Bulk or Segment — depending on how many people you're messaging.",
            "Type your message. Watch the character counter so you know how many SMS segments it costs.",
            "Add recipients. For Bulk, paste numbers one per line. For Segments, pick a saved audience.",
            "Choose Send Now or set a future date/time using the Schedule toggle.",
            "Review the cost summary, then click Send SMS.",
            "Watch the progress bar — it shows how many messages have been queued, sent and delivered.",
        ],
        benefits=[
            "Live cost calculation means no surprise bills — you see the total before you send.",
            "Scheduling lets you write at any hour but deliver at the best time.",
            "Bulk mode handles thousands of recipients without leaving your browser.",
        ],
        tip="A single SMS is 160 characters in plain text. Going over splits the message into multiple segments — each charged separately.",
    ),
    PageDoc(
        slug="campaigns",
        title="Campaigns",
        route="/campaigns",
        purpose="Design, schedule and monitor bulk messaging campaigns. Unlike a quick send, a campaign is tracked end-to-end with stats for sent, delivered and read.",
        sections=[
            "Search bar and filters for status and channel.",
            "Create Campaign button in the top right.",
            "Campaign list/cards showing Name, Status badge, Recipients, Sent/Delivered/Read counts, progress bar and estimated cost.",
            "Per-row action menu — Edit, Copy, Delete, Pause, Resume, View Details.",
            "Campaign details modal with full stats, message preview and recipient breakdown.",
        ],
        how_to=[
            "Click Create Campaign.",
            "Name the campaign so you can find it later (e.g. 'Eid Promo 2026').",
            "Pick a channel — SMS or WhatsApp.",
            "Choose the audience — all contacts, a tag, a saved segment, or a CSV upload.",
            "Write the message or pick a saved template.",
            "Choose Send Now or schedule for a future date and time.",
            "Review the cost summary and click Launch.",
            "Track progress on this page — Sent → Delivered → Read counts update live.",
        ],
        benefits=[
            "Campaigns give per-recipient delivery and read tracking, not just 'sent'.",
            "You can pause and resume mid-campaign if something goes wrong.",
            "Copy a previous campaign to launch a similar one in seconds.",
        ],
        tip="Always test a campaign on yourself first — pick the 'Send to one number' option, send to your own phone, then launch the full campaign once you're happy with the preview.",
    ),
    PageDoc(
        slug="templates",
        title="Templates",
        route="/templates",
        purpose="A library of reusable messages for SMS and WhatsApp. Templates save typing, keep your brand voice consistent, and (for WhatsApp) get pre-approved by Meta for higher delivery rates.",
        sections=[
            "Search bar and filters for channel, category and language.",
            "Template gallery showing Name, Category, Language, Channel badge, Approval status, Usage count and Last used date.",
            "Per-row actions — Use, Edit, Copy, Delete, Mark as Favorite.",
            "Create/Edit dialog with Name, Category, Language, Channel, Body text, Description and Variable placeholders.",
        ],
        how_to=[
            "Click Create Template.",
            "Give the template a clear name (e.g. 'Order Confirmation EN').",
            "Pick the channel (SMS or WhatsApp), language and category.",
            "Type the body, using {{name}} or {{order_id}} as variables that will be filled in per recipient.",
            "Save the template — it now appears in the gallery.",
            "When sending or running a campaign, choose 'Use Template' to pull it in.",
        ],
        benefits=[
            "Less typing, fewer mistakes, consistent brand voice.",
            "WhatsApp templates approved by Meta are required for outbound business-initiated chats.",
        ],
    ),
    PageDoc(
        slug="purchase-sms",
        title="Buy SMS Credits",
        route="/sms/purchase",
        purpose="Top up your SMS or WhatsApp credits using mobile money. Without credits you cannot send messages.",
        sections=[
            "Service toggle — SMS or WhatsApp.",
            "Predefined packages (Lite, Standard, Pro) with per-SMS price.",
            "Custom amount input — enter how much you want to spend and see how many credits you'll get.",
            "Payment method selector — mobile money providers supported in your country.",
            "User info form — Name, Email, Phone.",
            "Current balance card showing what you have right now.",
            "Invoice preview before payment is confirmed.",
        ],
        how_to=[
            "Choose SMS or WhatsApp at the top.",
            "Pick a package, or type a custom amount.",
            "Confirm the cost shown and the credits you will receive.",
            "Pick a payment method (e.g. M-Pesa, Airtel Money).",
            "Enter the phone number registered to your mobile money account.",
            "Click Pay and approve the prompt on your phone.",
            "Credits land in your account within seconds of payment confirmation.",
        ],
        benefits=[
            "Pay only for what you use — no monthly subscription needed.",
            "Larger packages get a lower per-SMS price.",
            "Mobile money payments are confirmed instantly.",
        ],
        tip="Always top up before launching a big campaign — campaigns pause if your balance hits zero mid-send.",
    ),
    PageDoc(
        slug="sender-names",
        title="Sender IDs (Branded Sender Names)",
        route="/sms/sender-names",
        purpose=(
            "A Sender ID is the name your recipients see at the top of every SMS — e.g. 'YOURBRAND' instead of an unknown phone number. "
            "Branded sender IDs build trust, get higher open rates, and protect your customers from impersonators. "
            "Regulators require approval before a sender ID can be used, so every request goes through a short review and payment step."
        ),
        sections=[
            "Sender ID request form — Requested ID (the name), Purpose, Sample message content, supporting attachments.",
            "Active Sender IDs table with a Status badge: Pending → Verifying → Awaiting Payment → Approved (or Rejected).",
            "Per-row actions — Set as Default, Edit, Delete, Request Payment.",
            "Request history showing every submission and its outcome.",
            "Link to upload KYC documents required by the regulator.",
        ],
        how_to=[
            "Click 'Request New Sender ID'.",
            "Type the sender name exactly as you want recipients to see it (max 11 characters, letters and numbers only, no spaces). Example: 'MIFUMOLABS'.",
            "Explain the purpose — what messages will be sent under this name (e.g. 'order confirmations, marketing'). Be honest, the regulator reads this.",
            "Paste a sample message so the regulator knows what kind of content will go out.",
            "Attach the KYC documents (business registration, ID of the authorised signatory).",
            "Submit the request — the status becomes Pending.",
            "Once the regulator's team verifies the request, the status changes to 'Awaiting Payment'.",
            "Click 'Request Payment' on that row — you will be redirected to the payment page.",
            "Complete payment via mobile money. The sender ID is then activated and the status turns Approved.",
            "Open Send SMS or Campaigns and choose the new sender ID from the dropdown.",
        ],
        benefits=[
            "Branded SMS — recipients see your business name, not a random number. This dramatically increases open rates.",
            "Trust and anti-fraud — customers can spot fake messages pretending to be you.",
            "Bonus SMS credits — once your sender ID is paid for and approved, SENDA automatically credits a bonus pack of SMS to your wallet as a thank-you.",
            "One business can register multiple sender IDs (e.g. separate IDs for marketing vs. transactional traffic).",
        ],
        tip=(
            "Pay only after the status changes to 'Awaiting Payment' — that means the regulator has accepted the name. "
            "Paying before approval risks losing the fee if the requested name is denied."
        ),
    ),
    PageDoc(
        slug="purchase-history",
        title="Purchase History",
        route="/sms/purchase-history",
        purpose="A full record of every credit purchase and sender ID payment — for accounting, audit and invoice download.",
        sections=[
            "Search bar and status filter (All / Completed / Pending / Failed).",
            "Transaction table — Invoice number, Package, Credits, Amount, Payment method, Status badge, Date.",
            "Action menu per row — View Details, Download Invoice.",
            "Detail sheet showing the full receipt and payment reference.",
        ],
        how_to=[
            "Open Purchase History from the SMS or Settings menu.",
            "Filter by status to find a specific transaction (e.g. failed payments worth retrying).",
            "Click any row to see the full receipt.",
            "Click Download Invoice to save a PDF for your accountant.",
        ],
        benefits=[
            "All your billing history in one place — no need to dig through emails.",
            "PDF invoices are ready for tax filing.",
        ],
    ),
    PageDoc(
        slug="sender-id-kyc",
        title="Sender ID KYC Upload",
        route="(opened from Sender IDs)",
        purpose="The page where you upload the legal documents required to approve a branded sender ID.",
        sections=[
            "Info cards — required documents, expected processing time, security/privacy statement.",
            "Document type selector (Business registration, Director ID, Authorisation letter).",
            "PDF file upload with size limit.",
            "Confirmation checkbox confirming the documents are genuine.",
            "Success state with a reference number to quote in support requests.",
        ],
        how_to=[
            "Have the original PDF documents ready on your device.",
            "Pick the document type from the dropdown.",
            "Click Upload and choose the file (PDF, under 5 MB).",
            "Tick the confirmation checkbox.",
            "Click Submit — note the reference number shown on the success screen.",
        ],
        tip="Documents are encrypted at rest and only accessible to the verification team. They are never shared with third parties.",
    ),
    PageDoc(
        slug="delivery-reports",
        title="Delivery Reports",
        route="(opened from SMS)",
        purpose="See exactly which messages were delivered, which are still pending, and which failed — per recipient.",
        sections=[
            "Date range and status filters.",
            "Search bar to find a specific number or message.",
            "Report table — Message ID, Status badge, Recipient, content preview, Sender ID, timestamp.",
            "Download button to export the report as CSV or PDF.",
            "Pagination controls for large reports.",
        ],
        how_to=[
            "Pick the date range you want to inspect.",
            "Filter by status (e.g. show only Failed to see what needs retry).",
            "Use the search bar to look up one specific recipient.",
            "Click Download to export the report.",
        ],
        benefits=[
            "Per-recipient visibility — you know precisely who got the message.",
            "Failed delivery details help you clean your contact list.",
        ],
    ),
    PageDoc(
        slug="whatsapp-cloud",
        title="WhatsApp",
        route="/whatsapp",
        purpose="The WhatsApp Cloud API hub. Send WhatsApp messages, manage templates, run polls and broadcast media to thousands of contacts.",
        sections=[
            "Three tabs — Quick Send, Bulk Send, Broadcast.",
            "Message composer with template selector and per-variable inputs.",
            "Recipient input — single number, bulk upload, contact group.",
            "Template management — Create, Edit, Delete.",
            "Poll creator for interactive surveys.",
            "Image / video / document upload for media broadcasts.",
            "Cost calculator and live send progress indicator.",
        ],
        how_to=[
            "For a one-off message — open Quick Send, pick a template, fill in the variables, type the recipient number, click Send.",
            "For a campaign — open Bulk Send, pick a template, upload a CSV with phone numbers and variables, click Send.",
            "For a poll — switch to Broadcast, click Create Poll, add the question and options, choose the audience and send.",
            "Track every send in the progress panel — messages flow from Queued → Sent → Delivered → Read.",
        ],
        benefits=[
            "Two-way conversations and rich media (images, PDFs, locations).",
            "Read receipts — you know when a recipient has actually opened the message.",
            "Cheaper than SMS for long messages because WhatsApp does not charge per 160 characters.",
        ],
        tip="Business-initiated WhatsApp messages must use a Meta-approved template — see the Create WhatsApp Template page next.",
    ),
    PageDoc(
        slug="create-whatsapp-template",
        title="Create WhatsApp Template",
        route="/whatsapp/templates/new",
        purpose="Build a WhatsApp Business template (text, media or interactive buttons) and submit it to Meta for approval. Approved templates are the only way to start a new WhatsApp conversation with a customer.",
        sections=[
            "Template name (lowercase letters, numbers and underscores only).",
            "Language selector — English or Kiswahili.",
            "Category selector — Marketing, Transactional, or OTP.",
            "Header type — None, Text, Image, Video or Document.",
            "Body text area with {{1}}, {{2}} placeholders for variables.",
            "Optional footer text.",
            "Buttons — Call to Action, Quick Reply, or URL.",
            "Live preview panel showing how the template will look on the recipient's phone.",
        ],
        how_to=[
            "Pick a descriptive name (e.g. 'order_confirmation_en').",
            "Choose the language and category that matches the use case.",
            "Pick the header type — for marketing, an image works well; for transactional, plain text is fine.",
            "Write the body. Use {{1}}, {{2}} to mark spots that will be filled in per recipient (like customer name or order ID).",
            "Add an optional footer (e.g. 'Reply STOP to opt out').",
            "Add buttons if useful — a URL button to your website, or a Quick Reply.",
            "Check the preview, then click Create Template — it is submitted to Meta for approval (usually under 24 hours).",
        ],
        tip="Marketing templates are reviewed more strictly than transactional ones. Pick the right category to avoid rejection.",
    ),
    PageDoc(
        slug="whatsapp-broadcast",
        title="WhatsApp Broadcast",
        route="/whatsapp-broadcast",
        purpose="Send the same WhatsApp message to many recipients at once. Best for promotions, reminders, and announcements.",
        sections=[
            "WhatsApp account selector (if multiple are connected).",
            "Message type — Text, Template or Image.",
            "Template picker when in Template mode.",
            "Recipient input — CSV upload, contact group, or paste numbers.",
            "Schedule options (Now or future date/time).",
            "Cost summary and a progress indicator during the broadcast.",
            "Success/fail result banner once the broadcast completes.",
        ],
        how_to=[
            "Pick the WhatsApp account that will send the broadcast.",
            "Choose the message type — for first contact with a recipient, you must pick Template.",
            "Select the template, upload a CSV with recipients and their variable values.",
            "Schedule for now or a future time.",
            "Review the cost summary and click Send Broadcast.",
            "Watch the result banner — it tells you how many were sent and how many failed (with reasons).",
        ],
    ),
    PageDoc(
        slug="ai-agents",
        title="AI Copilots",
        route="/ai-copilots",
        purpose="Build conversational AI agents — chatbots that answer customer questions, qualify leads, or run support 24/7 — without writing code.",
        sections=[
            "Agent setup wizard — Name, Industry, Language, Tone, Channel, Intent.",
            "Visual flow designer for building branching conversations.",
            "Node types — Message, List Menu, URL/CTA, Webhook.",
            "API endpoint configuration so the bot can talk to your own systems.",
            "Test / Preview mode that lets you chat with the bot before going live.",
            "Deploy button that pushes the bot to your chosen channel.",
        ],
        how_to=[
            "Click Create Agent and walk through the wizard (Name, Industry, Language, Tone).",
            "Pick the channel where the bot will run (WhatsApp or SMS).",
            "Drag nodes onto the canvas to design the conversation — Message for replies, List Menu for choices, Webhook to call your own API.",
            "Connect nodes with arrows to define the flow.",
            "Use Test mode to chat with the bot and check that replies are correct.",
            "Click Deploy when you're happy — the bot starts handling real conversations.",
        ],
        benefits=[
            "Handle thousands of customer messages without growing the support team.",
            "Always-on — bot answers at 3am the same as 3pm.",
            "Integrates with your existing systems via webhooks.",
        ],
    ),
    PageDoc(
        slug="voice-agents",
        title="Voice Copilots (Private Beta)",
        route="/voice-copilots",
        purpose="Sign-up page for AI-driven voice copilots — bots that make and answer phone calls, transcribe conversations, and run interactive IVR menus. Currently in private beta.",
        sections=[
            "Feature grid — Voice Dashboard, Call Flows, Voice Campaigns, Call Logs, AI Voice Config, Create Copilot.",
            "Beta signup form — Name, Email, Phone, with optional KYC file upload.",
            "Perks list — free access during beta, direct line to the product team, ability to shape features before launch.",
            "Success confirmation after signup.",
        ],
        how_to=[
            "Read through the feature list to understand what's coming.",
            "Fill in your details on the signup form.",
            "Upload basic KYC if requested.",
            "Submit — the product team will reach out within a few days to onboard you.",
        ],
    ),
    PageDoc(
        slug="settings",
        title="Settings",
        route="/settings",
        purpose="Manage everything about your account — profile, security, billing, team, API keys and webhook integrations — from one place.",
        sections=[
            "Profile tab — avatar, name, email, phone, company, timezone.",
            "Security tab — 2FA setup via QR code, active device sessions, sign out everywhere.",
            "Billing tab — payment methods, subscription, invoice history.",
            "Team tab — invite members, change roles, remove users.",
            "API Keys tab — create, copy and revoke API tokens.",
            "Webhooks tab — create, edit and delete webhook endpoints.",
            "Integrations tab — connect third-party tools.",
        ],
        how_to=[
            "Open Settings from the sidebar or your avatar menu.",
            "Click the tab for what you want to change (Profile, Security, etc.).",
            "Make the changes and click Save.",
            "For Security, click Enable 2FA and scan the QR code with Google Authenticator or Authy.",
            "For Team, click Invite Member, enter the colleague's email and pick a role.",
            "For API Keys, click Create Key, name it (e.g. 'production-server'), copy the token and store it safely — it is shown only once.",
        ],
        tip="Enable 2FA for every account. It is the single biggest thing you can do to protect against unauthorised access.",
    ),
    PageDoc(
        slug="developer",
        title="Developer",
        route="/developer",
        purpose="Public developer landing page that showcases the SENDA API — capabilities, SDKs, code samples — for technical teams considering an integration.",
        sections=[
            "Hero section with a code example you can copy.",
            "Feature cards — SMS API, WhatsApp API, Analytics API, Webhooks, Rate limiting, Authentication.",
            "Quick-start guide.",
            "Multi-language code examples — Python, Node.js, PHP, cURL.",
            "SDK links and API pricing table.",
            "CTAs — View Full API Docs, Sign Up, Contact Sales.",
        ],
        how_to=[
            "Browse the features cards to confirm the API supports your use case.",
            "Copy the quick-start code and try it locally.",
            "Click View API Docs for the full reference.",
            "If you need volume pricing or enterprise features, click Contact Sales.",
        ],
    ),
    PageDoc(
        slug="integration-guide",
        title="Integration Guide (API Reference)",
        route="/integration-guide",
        purpose="The full REST API reference for developers integrating SENDA into their own apps or back-end systems.",
        sections=[
            "Table of contents sidebar.",
            "Endpoint cards — HTTP method, path, description, parameters, request body and example response.",
            "Code examples in Python, Node.js, PHP and cURL.",
            "Error code reference.",
            "Rate limiting and authentication overview.",
        ],
        how_to=[
            "Find the endpoint you need from the sidebar.",
            "Read the description and parameter table.",
            "Copy the example request and adjust it for your environment.",
            "Replace YOUR_API_KEY with a real key from Settings → API Keys.",
            "Test the call and check the response against the example.",
        ],
        tip="All endpoints require a valid Bearer token. Keep API keys out of client-side code — call SENDA from your back-end only.",
    ),
    PageDoc(
        slug="partner-integration",
        title="Partner Integration (Pertina API)",
        route="/partner-integration",
        purpose=(
            "Special API endpoints designed for Mifumo Connect Partners (resellers and integrators). "
            "Partners can query the credit balance and SMS usage of every tenant they manage and bill their own customers programmatically."
        ),
        sections=[
            "Four endpoint cards — Tenant Balance (GET), All Clients Balance (GET), Tenant Usage (GET), Global Usage (GET).",
            "Each card shows the HTTP method, full path, required parameters and a sample JSON response.",
            "Authentication and rate-limit notes specific to partner accounts.",
            "Copy-cURL buttons for quick testing.",
        ],
        how_to=[
            "Apply for partner status (see the 'Becoming a Partner' section later in this guide).",
            "Once approved, open Settings → API Keys and generate a key labelled 'partner-integration'.",
            "Copy the cURL example from any endpoint card and paste it into your terminal.",
            "Replace YOUR_API_KEY and TENANT_ID with real values.",
            "Use the returned JSON to power your own billing or reseller dashboard.",
        ],
        benefits=[
            "Resell SENDA under your own brand — your customers never know you use Mifumo behind the scenes.",
            "Programmatic balance and usage queries mean automatic invoicing.",
        ],
    ),
    PageDoc(
        slug="terms",
        title="Terms of Service",
        route="/terms",
        purpose="The legal terms and conditions governing use of the SENDA platform. Every user accepts these at signup.",
        sections=[
            "Service description.",
            "Acceptable use policy.",
            "User rights and responsibilities.",
            "Limitations of liability.",
            "Indemnification and termination clauses.",
            "Amendments and governing law.",
        ],
        how_to=[
            "Read the full document carefully when you sign up.",
            "Check back if you receive an email notifying you of an update.",
            "Contact the legal team via the link at the bottom if anything is unclear.",
        ],
    ),
    PageDoc(
        slug="privacy",
        title="Privacy Policy",
        route="/privacy",
        purpose="Explains how SENDA collects, uses, retains and protects user and contact data, in line with GDPR-style and local data-protection rules.",
        sections=[
            "Data we collect.",
            "How that data is used.",
            "Data retention periods.",
            "Your rights as a user — access, correction, deletion.",
            "Cookies and tracking.",
            "Third-party sharing (kept to a minimum).",
            "Security measures.",
            "Contact information for privacy questions.",
        ],
        how_to=[
            "Read it once at signup.",
            "Email privacy@mifumolabs.com if you want to exercise a data right (e.g. data export or deletion).",
        ],
    ),
    PageDoc(
        slug="not-found",
        title="404 Not Found",
        route="*",
        purpose="A friendly error page that appears when a user types a URL that does not exist — for example a stale bookmark or a typo.",
        sections=[
            "Large '404' heading.",
            "'Oops! Page not found' message.",
            "Return to Home button.",
        ],
        how_to=[
            "Click Return to Home to get back to the landing page.",
            "Use the browser's Back button if you arrived here by accident.",
        ],
    ),
]

# ---- low-level helpers --------------------------------------------------

def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_borders(cell, color_hex: str = SECTION_RULE) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color_hex)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _h1(doc: Document, text: str, color: RGBColor = BRAND_PRIMARY, size: int = 20) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BRAND_ACCENT


def _para(doc: Document, text: str, size: int = 10, color: RGBColor = TEXT_DARK,
          italic: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_DARK


def _numbered(doc: Document, items: List[str]) -> None:
    for i, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.4)
        num = p.add_run(f"{i}.  ")
        num.bold = True
        num.font.size = Pt(10)
        num.font.color.rgb = BRAND_PRIMARY
        body = p.add_run(item)
        body.font.size = Pt(10)
        body.font.color.rgb = TEXT_DARK


# ---- in-cell helpers (used inside the two-column table) ----------------

def _cell_eyebrow(cell, text: str, color: RGBColor) -> None:
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = color


def _cell_title(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BRAND_PRIMARY


def _cell_route(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_MUTED


def _cell_section_label(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = BRAND_ACCENT


def _cell_paragraph(cell, text: str, size: float = 10) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = TEXT_DARK


def _cell_bullets(cell, items: List[str]) -> None:
    for item in items:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.space_after = Pt(1)
        bullet = p.add_run("•  ")
        bullet.font.size = Pt(10)
        bullet.font.color.rgb = BRAND_PRIMARY
        body = p.add_run(item)
        body.font.size = Pt(9.5)
        body.font.color.rgb = TEXT_DARK


def _cell_numbered(cell, items: List[str]) -> None:
    for i, item in enumerate(items, start=1):
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.space_after = Pt(1)
        num = p.add_run(f"{i}.  ")
        num.bold = True
        num.font.size = Pt(10)
        num.font.color.rgb = BRAND_PRIMARY
        body = p.add_run(item)
        body.font.size = Pt(9.5)
        body.font.color.rgb = TEXT_DARK


def _cell_tip(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    label = p.add_run("Tip — ")
    label.bold = True
    label.italic = True
    label.font.size = Pt(9.5)
    label.font.color.rgb = BRAND_ACCENT
    body = p.add_run(text)
    body.italic = True
    body.font.size = Pt(9.5)
    body.font.color.rgb = TEXT_MUTED


def _fill_image_cell(cell, image_path: Path) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(6.2))


def _fill_text_cell(cell, page: PageDoc, display_index: int) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_borders(cell)

    _cell_eyebrow(cell, f"PAGE {display_index:02d}   ·   {page.route}", BRAND_ACCENT)
    _cell_title(cell, page.title)

    _cell_paragraph(cell, page.purpose)

    if page.sections:
        _cell_section_label(cell, "What's on this page")
        _cell_bullets(cell, page.sections)

    if page.how_to:
        _cell_section_label(cell, "How to use it")
        _cell_numbered(cell, page.how_to)

    if page.benefits:
        _cell_section_label(cell, "Why it matters")
        _cell_bullets(cell, page.benefits)

    if page.tip:
        _cell_tip(cell, page.tip)


# ---- composed sections --------------------------------------------------

def _add_cover(doc: Document, included: int) -> None:
    # Slim, professional cover — no oversized title.
    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eyebrow.paragraph_format.space_before = Pt(90)
    er = eyebrow.add_run("MIFUMO LABS")
    er.bold = True
    er.font.size = Pt(11)
    er.font.color.rgb = BRAND_ACCENT

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    tr = title.add_run("SENDA")
    tr.bold = True
    tr.font.size = Pt(34)
    tr.font.color.rgb = BRAND_PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("User Guide")
    sr.font.size = Pt(16)
    sr.font.color.rgb = TEXT_DARK

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.paragraph_format.space_before = Pt(6)
    tagr = tag.add_run("A page-by-page walkthrough of the SENDA platform")
    tagr.italic = True
    tagr.font.size = Pt(11)
    tagr.font.color.rgb = TEXT_MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(60)
    mr = meta.add_run(f"{included} screens documented   ·   Mifumo Labs   ·   Edition 1")
    mr.font.size = Pt(10)
    mr.font.color.rgb = TEXT_FAINT

    _add_page_break(doc)


def _add_acknowledgement(doc: Document) -> None:
    _h1(doc, "Acknowledgement")
    _para(
        doc,
        "This guide is the work of the Mifumo Labs product and engineering team. "
        "We thank every business, partner and individual who has tested SENDA, reported issues, "
        "and shared the kind of honest feedback that makes a product better — your ideas are visible on every screen in this book.",
    )
    _para(
        doc,
        "We also acknowledge our customers across Tanzania, Kenya, Uganda and the wider region who trust SENDA with their daily communications. "
        "Building reliable messaging infrastructure for African businesses is the reason this platform exists.",
    )
    _para(doc, "— The SENDA Team", italic=True, color=TEXT_MUTED)
    _add_page_break(doc)


def _add_about_us(doc: Document) -> None:
    _h1(doc, "About SENDA")
    _para(
        doc,
        "SENDA is a unified business messaging platform built by Mifumo Labs. It brings SMS, WhatsApp Cloud API, AI Copilots and voice "
        "calling together in one dashboard so that any business — from a small shop to a national enterprise — can talk to its customers reliably.",
    )

    _h2(doc, "What you can do with SENDA")
    _bullet(doc, "Send branded SMS to one customer or one million.")
    _bullet(doc, "Run WhatsApp Business campaigns with templates, media and polls.")
    _bullet(doc, "Build AI chatbots that answer questions 24/7.")
    _bullet(doc, "Manage contacts, tags and segments in one shared database.")
    _bullet(doc, "Track every message — sent, delivered, read — with per-recipient detail.")
    _bullet(doc, "Top up credits, request branded sender IDs, and pay by mobile money.")

    _h2(doc, "Why we built it")
    _para(
        doc,
        "Most messaging tools were built for North America or Europe and ignore the realities of African networks, regulators and payment methods. "
        "SENDA was built locally, in collaboration with the businesses that use it, with mobile money payments, regional sender-ID compliance "
        "and Kiswahili language support as first-class features — not afterthoughts.",
    )

    _h2(doc, "Who this guide is for")
    _para(
        doc,
        "Anyone who uses SENDA in their day-to-day work — business owners, marketers, support agents, developers and partners. "
        "Each section pairs a real screenshot from the mobile experience with a short, plain-language description so you can recognise every "
        "page at a glance and learn what to do on it.",
    )
    _add_page_break(doc)


def _add_toc(doc: Document, included: List[Tuple[int, PageDoc]]) -> None:
    _h1(doc, "Table of Contents")
    _para(
        doc,
        "Each page below is documented with a screenshot and a short walkthrough. "
        "Pages that are coming soon (not yet shipped) are not included in this edition.",
        italic=True,
        color=TEXT_MUTED,
    )
    # Front-matter entries
    for label in [
        "Acknowledgement",
        "About SENDA",
        "Page-by-page walkthrough",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"  •  {label}")
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT_DARK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

    _h2(doc, "Pages")
    for display_index, page in included:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        n = p.add_run(f"{display_index:02d}.  ")
        n.bold = True
        n.font.size = Pt(10.5)
        n.font.color.rgb = BRAND_ACCENT
        t = p.add_run(page.title)
        t.font.size = Pt(10.5)
        t.font.color.rgb = TEXT_DARK
        r = p.add_run(f"   {page.route}")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = TEXT_MUTED

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)

    for label in ["Becoming a Partner", "The Team"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"  •  {label}")
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT_DARK

    _add_page_break(doc)


def _add_page_section(doc: Document, page: PageDoc, image_path: Path, display_index: int) -> None:
    image_on_left = (display_index % 2 == 1)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    left, right = table.rows[0].cells
    left.width = Cm(7.5) if image_on_left else Cm(9.5)
    right.width = Cm(9.5) if image_on_left else Cm(7.5)

    if image_on_left:
        _fill_image_cell(left, image_path)
        _fill_text_cell(right, page, display_index)
    else:
        _fill_text_cell(left, page, display_index)
        _fill_image_cell(right, image_path)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    _add_page_break(doc)


def _add_partner_program(doc: Document) -> None:
    _h1(doc, "Becoming a Partner")
    _para(
        doc,
        "The Mifumo Connect Partner Program is for resellers, agencies and integrators who want to offer SENDA to their own customers — "
        "either under their own brand or as a value-added service.",
    )

    _h2(doc, "What partners get")
    _bullet(doc, "Access to the Partner API — query the balance and SMS usage of every tenant you manage.")
    _bullet(doc, "The Partner Insights dashboard — analytics across all your client accounts in one view.")
    _bullet(doc, "Bulk pricing — better per-SMS rates that you can re-sell at your own margin.")
    _bullet(doc, "A dedicated account manager and priority support.")
    _bullet(doc, "Optional white-label branding so end users see your business, not Mifumo.")

    _h2(doc, "How to apply")
    _numbered(
        doc,
        [
            "Sign up for a regular SENDA account if you don't already have one.",
            "Send an email to partners@mifumolabs.com with your company name, target market and expected monthly volume.",
            "Provide basic KYC — business registration document and the ID of an authorised signatory.",
            "Sign the partnership agreement we send back.",
            "Once approved, the team enables 'Partner' status on your account — the Partner Integration and Partner Insights routes unlock automatically.",
            "Generate a partner API key under Settings → API Keys and start integrating.",
        ],
    )

    _para(
        doc,
        "Tip — partners with consistent monthly volume qualify for revenue-share on every new customer they bring onto SENDA.",
        italic=True,
        color=TEXT_MUTED,
    )
    _add_page_break(doc)


def _add_team(doc: Document) -> None:
    _h1(doc, "The Team")
    _para(
        doc,
        "SENDA is built by a small, focused team based in East Africa. "
        "Every line of code, every screen in this guide, and every customer conversation comes from the three people below.",
    )

    members = [
        ("Mgasa Lusas", "Chief Executive Officer (CEO)",
         "Sets the vision, leads partnerships and makes sure SENDA stays focused on what actually helps businesses grow."),
        ("Florence Sway", "Software Developer",
         "Owns the user experience — designs and builds the dashboards, mobile flows and everything customers see on screen."),
        ("Magessa", "Backend Developer",
         "Builds and runs the infrastructure that delivers every SMS and WhatsApp message — the engine room of the platform."),
    ]

    for name, role, bio in members:
        # Card-ish layout: bold name, accent role, body bio
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(0)
        name_run = p.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(13)
        name_run.font.color.rgb = BRAND_PRIMARY

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        role_run = p2.add_run(role)
        role_run.bold = True
        role_run.font.size = Pt(10.5)
        role_run.font.color.rgb = BRAND_ACCENT

        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(4)
        bio_run = p3.add_run(bio)
        bio_run.font.size = Pt(10)
        bio_run.font.color.rgb = TEXT_DARK

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    close = p.add_run("Thank you for using SENDA.")
    close.italic = True
    close.font.size = Pt(11)
    close.font.color.rgb = TEXT_MUTED


# ---- main ---------------------------------------------------------------

def _resolve_image(page: PageDoc) -> Optional[Path]:
    # Match by slug, ignoring whatever numeric prefix the file uses.
    # That way pages can be added, removed or reordered without renaming files.
    for ext in ("png", "jpg", "jpeg"):
        matches = sorted(SCREENSHOT_DIR.glob(f"*-{page.slug}.{ext}"))
        if matches:
            return matches[0]
        # Also accept the slug as a bare filename, e.g. "dashboard.png".
        bare = SCREENSHOT_DIR / f"{page.slug}.{ext}"
        if bare.exists():
            return bare
    return None


def build() -> Tuple[int, int]:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Resolve which pages have an image — that's the only way a page makes it in.
    included: List[Tuple[int, PageDoc, Path]] = []  # (display_index, page, image_path)
    for page in PAGES:
        img = _resolve_image(page)
        if img:
            display_index = len(included) + 1
            included.append((display_index, page, img))

    total = len(PAGES)
    shipped = len(included)

    # Cover + front matter
    _add_cover(doc, shipped)
    _add_acknowledgement(doc)
    _add_about_us(doc)
    _add_toc(doc, [(i, p) for (i, p, _) in included])

    # Walkthrough
    _h1(doc, "Page-by-page walkthrough")
    _para(
        doc,
        "Mobile screenshots are paired with descriptions, alternating sides on every page so the eye keeps moving. "
        "Each entry explains what the page is, what's on it, how to use it and why it matters.",
        italic=True,
        color=TEXT_MUTED,
    )
    _add_page_break(doc)

    for display_index, page, image_path in included:
        _add_page_section(doc, page, image_path, display_index)

    _add_partner_program(doc)
    _add_team(doc)

    doc.save(OUTPUT_PATH)
    return shipped, total


if __name__ == "__main__":
    shipped, total = build()
    print(f"Generated: {OUTPUT_PATH}")
    print(f"Pages included: {shipped} of {total} (others have no screenshot yet — treated as coming soon).")
